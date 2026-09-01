from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MANUSCRIPT_ANONYMOUS.md"
OUT = ROOT / "MANUSCRIPT_ANONYMOUS_v2.docx"
BODY_FONT = "Times New Roman"
BODY_SIZE = 10.5
ARABIC_RE = re.compile(r"[\u0600-\u06FF]")


def set_run_font(run, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = BODY_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_rtl(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))
    for run in paragraph.runs:
        rpr = run._r.get_or_add_rPr()
        if rpr.find(qn("w:rtl")) is None:
            rpr.append(OxmlElement("w:rtl"))


def add_inline(paragraph, text, size=BODY_SIZE):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            set_run_font(paragraph.add_run(text[pos:m.start()]), size)
        token = m.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size, bold=True)
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Menlo"
            r.font.size = Pt(max(8, size - 1))
        else:
            set_run_font(paragraph.add_run(token[1:-1]), size, italic=True)
        pos = m.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), size)
    if ARABIC_RE.search(text):
        set_rtl(paragraph)


def format_paragraph(p, before=0, after=5, line=1.06, left=0, first=0):
    f = p.paragraph_format
    f.space_before = Pt(before)
    f.space_after = Pt(after)
    f.line_spacing = line
    if left:
        f.left_indent = Inches(left)
    if first:
        f.first_line_indent = Inches(first)


def add_body(doc, text):
    p = doc.add_paragraph()
    add_inline(p, text)
    format_paragraph(p)
    return p


def set_cell_margins(cell, top=40, start=50, bottom=40, end=50):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}")) or OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            tcMar.append(node)


def add_table(doc, rows):
    header = rows[0]
    data = [r for r in rows[1:] if not all(re.fullmatch(r"\s*:?-+:?\s*", c or "") for c in r)]
    ncols = max(len(r) for r in [header] + data)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout"); tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    font_size = 7.0 if ncols >= 7 else 8.0 if ncols >= 5 else 9.0
    widths = [6.75 / ncols] * ncols
    if ncols >= 5:
        widths[0] = 1.60
        rem = 6.75 - widths[0]
        widths[1:] = [rem / (ncols - 1)] * (ncols - 1)
    elif ncols == 3:
        widths = [2.9, 1.9, 1.95]
    elif ncols == 2:
        widths = [4.7, 2.05]
    for c in range(ncols):
        table.columns[c].width = Inches(widths[c])
    for c, text in enumerate(header):
        cell = table.rows[0].cells[c]
        cell.width = Inches(widths[c])
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, text.strip(), font_size)
        for run in p.runs:
            run.bold = True
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))
    for row in data:
        cells = table.add_row().cells
        for c in range(ncols):
            cell = cells[c]
            cell.width = Inches(widths[c])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            text = row[c].strip() if c < len(row) else ""
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, text, font_size)
            p.paragraph_format.space_after = Pt(0)
        cant = OxmlElement("w:cantSplit")
        table.rows[-1]._tr.get_or_add_trPr().append(cant)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        rows.append([c.strip() for c in raw.split("|")])
        i += 1
    return rows, i


def build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72); sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82); sec.right_margin = Inches(0.82)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT; normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""
    seen_section = False
    in_references = False
    i = 0
    while i < len(lines):
        s = lines[i].rstrip()
        if not s.strip() or s.strip() == "---":
            i += 1; continue
        if s.strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows); continue
        if s.startswith("# "):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, s[2:].strip(), 16)
            for r in p.runs: r.bold = True
            format_paragraph(p, after=6, line=1.0)
        elif s.startswith("## "):
            seen_section = True
            heading_text = s[3:].strip()
            if heading_text == "13. References":
                in_references = True
            p = doc.add_paragraph()
            add_inline(p, heading_text, 13)
            for r in p.runs: r.bold = True
            format_paragraph(p, before=9, after=4, line=1.0)
        elif s.startswith("### "):
            p = doc.add_paragraph()
            if not seen_section:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(p, s[4:].strip(), 11)
                for r in p.runs: r.bold = True
                format_paragraph(p, after=3, line=1.0)
            else:
                add_inline(p, s[4:].strip(), 11)
                for r in p.runs: r.bold = True
                format_paragraph(p, before=6, after=3, line=1.0)
        elif re.match(r"^\s*[-*]\s+", s):
            parts = [re.sub(r"^\s*[-*]\s+", "", s).strip()]
            j = i + 1
            while j < len(lines):
                nxt = lines[j].rstrip()
                if (not nxt.strip() or nxt.startswith("#") or nxt.strip().startswith("|")
                    or nxt.startswith("> ") or re.match(r"^\s*[-*]\s+", nxt)
                    or re.match(r"^\s*\d+\.\s+", nxt) or nxt.strip() == "---"):
                    break
                parts.append(nxt.strip()); j += 1
            p = doc.add_paragraph(); p.style = doc.styles["List Bullet"]
            add_inline(p, " ".join(parts))
            format_paragraph(p, after=2, line=1.04)
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            i = j - 1
        elif re.match(r"^\s*\d+\.\s+", s):
            m = re.match(r"^\s*(\d+)\.\s+(.*)$", s)
            parts = [m.group(2).strip()]
            j = i + 1
            while j < len(lines):
                nxt = lines[j].rstrip()
                if (not nxt.strip() or nxt.startswith("#") or nxt.strip().startswith("|")
                    or nxt.startswith("> ") or re.match(r"^\s*[-*]\s+", nxt)
                    or re.match(r"^\s*\d+\.\s+", nxt) or nxt.strip() == "---"):
                    break
                parts.append(nxt.strip()); j += 1
            p = doc.add_paragraph()
            add_inline(p, f"{m.group(1)}. " + " ".join(parts))
            format_paragraph(p, after=2, line=1.04)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            i = j - 1
        elif s.startswith("> "):
            parts = [s[2:].strip()]
            j = i + 1
            while j < len(lines) and lines[j].rstrip().startswith("> "):
                parts.append(lines[j].rstrip()[2:].strip()); j += 1
            p = doc.add_paragraph()
            add_inline(p, " ".join(parts), 10.2)
            format_paragraph(p, before=2, after=4, line=1.04, left=0.35)
            p.paragraph_format.right_indent = Inches(0.2)
            i = j - 1
        elif s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(s.strip("*")), 10.2, italic=True)
            format_paragraph(p, after=6, line=1.0)
        else:
            # join wrapped markdown lines into one paragraph until a structural boundary
            parts = [s.strip()]
            j = i + 1
            while j < len(lines):
                nxt = lines[j].rstrip()
                if (not nxt.strip() or nxt.startswith("#") or nxt.strip().startswith("|")
                    or nxt.startswith("> ") or re.match(r"^\s*[-*]\s+", nxt)
                    or re.match(r"^\s*\d+\.\s+", nxt) or nxt.strip() == "---"):
                    break
                parts.append(nxt.strip()); j += 1
            text = " ".join(parts)
            if in_references:
                p = doc.add_paragraph(); add_inline(p, text, 9.2)
                format_paragraph(p, after=2, line=1.0)
            else:
                p = add_body(doc, text)
            i = j - 1
        i += 1
    doc.save(OUT)
    print(f"WROTE {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
